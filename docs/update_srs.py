#!/usr/bin/env python3
"""Update Eunice's SRS document with shipped features and new feature ideas."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
BRAND = RGBColor(0x1D, 0x4E, 0x89)


def set_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = BRAND


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph("")
    return t


def bold_para(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    p.add_run(text)


def create_srs():
    doc = Document()
    set_style(doc)

    # ---- COVER PAGE ----
    for _ in range(5):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SOFTWARE REQUIREMENTS SPECIFICATION\n& BUSINESS PLAN")
    r.font.size = Pt(26)
    r.font.color.rgb = BRAND
    r.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Nyumba Yangu Online")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "A rental & property-management marketplace for Dar es Salaam — built so a tenant can\n"
        "find and inquire about a home without an account, and a landlord can run their whole\n"
        "portfolio (units, leases, rent, arrears) from one dashboard."
    )
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r3.italic = True

    doc.add_paragraph("")
    table(doc,
        ["Domain", "Version", "Audience"],
        [["nyumbayangu.online", "2.0 — updated August 2026", "Eunice Chihoma (CEO) & Erick Mkingule (CTO)"]]
    )

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(
        "Prepared by:\n"
        "Eunice Chihoma — Co-Founder & CEO\n"
        "Erick Mkingule — Co-Founder & CTO\n\n"
        "August 2026"
    )
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ---- STATUS KEY ----
    heading(doc, "Status Key")
    doc.add_paragraph(
        "SHIPPED — already built and working on web + mobile\n"
        "FRONTEND-ONLY — buildable without touching the backend\n"
        "NEEDS BACKEND — needs new API/DB work\n"
        "OPEN DECISION — needs both co-founders to agree first\n"
        "NEW — feature added in this version of the SRS"
    )
    doc.add_paragraph("")

    # ---- TABLE OF CONTENTS ----
    heading(doc, "Contents")
    toc = [
        "01 — WHY THIS EXISTS (Executive Summary)",
        "02 — WHO IT'S FOR (Users & Roles)",
        "03 — HOW IT MAKES MONEY (Business Model)",
        "04 — WHAT IT'S BUILT ON (System Architecture)",
        "05 — ROADMAP (Scope & Phasing)",
        "06 — REQUIREMENTS: Public & Tenant Portal",
        "07 — REQUIREMENTS: Landlord/Agent Portal",
        "08 — REQUIREMENTS: Admin Portal",
        "09 — REQUIREMENTS: Mobile App",
        "10 — DATA MODEL (Contracts)",
        "11 — API REFERENCE",
        "12 — NON-FUNCTIONAL REQUIREMENTS",
        "13 — CORRECTIONS TO ORIGINAL DRAFT",
        "14 — OPEN DECISIONS & BACKLOG",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # ---- 01 WHY THIS EXISTS ----
    heading(doc, "01 — WHY THIS EXISTS")
    heading(doc, "Executive Summary", level=2)
    doc.add_paragraph(
        "Renting in Dar es Salaam mostly runs through word of mouth and roadside signboards. "
        "Nyumba Yangu Online gives landlords and agents a place to list vacant units and manage the "
        "ones already occupied, and gives tenants a place to browse and contact a lister without "
        "needing an account first. The product has three audiences and, correspondingly, three portals: "
        "a public marketplace, a landlord/agent management dashboard, and a lightweight tenant portal "
        "for people already renting through the platform."
    )
    doc.add_paragraph(
        "The platform is live at nyumbayangu.online with both a web app and a native mobile app "
        "(iOS + Android via React Native / Expo). The backend serves 48+ API endpoints."
    )

    heading(doc, "The Problem", level=2)
    doc.add_paragraph(
        "Landlords track units, tenants, and rent in notebooks or memory. Renters can't search or "
        "compare listings anywhere — they rely on brokers (dalali) and referrals. Dalali charge 1-3 "
        "months' rent as fees. Tenants have no digital lease records, no payment receipts, and no way "
        "to report maintenance issues."
    )

    heading(doc, "The Wedge", level=2)
    doc.add_paragraph(
        "Free, searchable listings pull tenants in. Landlords get real management tools (occupancy, "
        "leases, arrears, amenities, maintenance) as the reason to stay, not just to list once."
    )

    heading(doc, "The Revenue", level=2)
    doc.add_paragraph(
        "Landlords pay a small per-unit monthly fee once they're managing units through the platform. "
        "Tenants never pay to browse, save, or inquire."
    )
    doc.add_page_break()

    # ---- 02 WHO IT'S FOR ----
    heading(doc, "02 — WHO IT'S FOR")
    heading(doc, "Users & Roles", level=2)
    doc.add_paragraph(
        "Four roles exist in the system today. A visitor needs no account at all for the core browsing "
        "experience — that's deliberate, and it's the first thing to protect when adding any new feature."
    )

    heading(doc, "Guest (Tenant / Buyer)", level=3)
    doc.add_paragraph(
        "Searches and filters listings, opens full detail pages, views amenities, and saves favorites "
        "— all without an account. Registering (as role tenant) unlocks a portal showing their active "
        "lease, rent-due tracking, payment history, and the ability to submit maintenance requests."
    )

    heading(doc, "Landlord / Agent", level=3)
    doc.add_paragraph(
        "Runs their portfolio from one dashboard. Lists properties and units (with full amenity details), "
        "publishes listings, tracks occupancy (including under-renovation status), records tenants and "
        "leases, logs rent payments, chases arrears, manages maintenance requests, and sees portfolio-wide "
        "reports. Pays a subscription once they're managing units, priced by how many units they hold."
    )

    heading(doc, "Admin", level=3)
    doc.add_paragraph(
        "Reviews landlord/agent KYC (NIDA number or business license) and property submissions, "
        "approving or rejecting each. This is what earns a listing its \"Verified\" badge in search results."
    )

    heading(doc, "System (background)", level=3)
    doc.add_paragraph(
        "Every authenticated call carries a 7-day JWT. Login is either phone + password, or a 6-digit "
        "one-time code emailed via Resend — no separate \"confirm your email\" step exists outside of "
        "that OTP flow."
    )
    doc.add_page_break()

    # ---- 03 HOW IT MAKES MONEY ----
    heading(doc, "03 — HOW IT MAKES MONEY")
    heading(doc, "Business Model & Monetization", level=2)
    doc.add_paragraph(
        "Tenants are free, permanently — that's the acquisition engine. Landlords are the paying side."
    )

    heading(doc, "Subscription Tiers", level=3)
    table(doc,
        ["Tier", "Units Held", "Price / Unit / Month", "Notes"],
        [
            ["Free", "1 unit", "TZS 0", "Seeds adoption — every landlord's first unit is free"],
            ["Starter", "2-5 units", "TZS 10,000", "Small portfolio landlords"],
            ["Growth", "6-20 units", "TZS 8,000", "Growing landlords"],
            ["Scale", "21+ units", "TZS 6,000", "Volume discount for larger portfolios / agencies"],
        ]
    )
    doc.add_paragraph(
        "Why per-unit, not per-listing: it rewards landlords for keeping their portfolio current in the "
        "system (occupied units still count) rather than only for the moment they publish a listing, "
        "which is what makes the management tools — not just the listing — the reason to keep paying."
    )

    heading(doc, "Future Revenue Streams", level=3)
    doc.add_paragraph(
        "- Featured listings: landlords pay to boost visibility (TZS 10,000-50,000 per listing)\n"
        "- Transaction fees: small percentage on M-Pesa rent payments (1-2%)\n"
        "- Tenant deposit insurance partnerships\n"
        "- Anonymized market data for developers and investors\n"
        "- Agent/property manager white-label accounts"
    )

    heading(doc, "Monetization Gap", level=3)
    doc.add_paragraph(
        "Tier and price recompute automatically from the landlord's live unit count — there is no "
        "plan-selection screen, and nothing is billed on a schedule yet. NEEDS BACKEND: a real "
        "collection mechanism (M-Pesa or invoice) is the biggest monetization gap."
    )
    doc.add_page_break()

    # ---- 04 WHAT IT'S BUILT ON ----
    heading(doc, "04 — WHAT IT'S BUILT ON")
    heading(doc, "System Architecture", level=2)
    doc.add_paragraph(
        "One repository, three apps: a Go REST API talking directly to Postgres with hand-written SQL, "
        "a React single-page web app, and a React Native mobile app (iOS + Android). The backend is "
        "the single source of truth for all clients."
    )

    heading(doc, "Backend", level=3)
    doc.add_paragraph(
        "- Go, routed with chi, one handler struct per domain\n"
        "- pgx/pgxpool straight to PostgreSQL — raw SQL, no ORM\n"
        "- JWT auth (7-day expiry) + bcrypt password hashing\n"
        "- 48+ API endpoints across public, tenant, landlord, and admin groups\n"
        "- Idempotent SQL migrations embedded in the binary\n"
        "- Deployed to a single VPS (Vultr, Nairobi data center)\n"
        "- Uploaded images live on that machine's local disk"
    )

    heading(doc, "Web Frontend", level=3)
    doc.add_paragraph(
        "- React + TypeScript (.tsx), built with Vite\n"
        "- Tailwind CSS, Lucide icons, React Router\n"
        "- Global state via React Context — auth, wishlist each get their own provider\n"
        "- API calls go through one hand-written fetch wrapper (no Axios)\n"
        "- Live at https://nyumbayangu.online"
    )

    heading(doc, "Mobile App (NEW)", level=3)
    doc.add_paragraph(
        "- React Native + Expo (TypeScript), SDK 57\n"
        "- React Navigation (native stack + bottom tabs)\n"
        "- AsyncStorage with in-memory cache (replacing localStorage)\n"
        "- @expo/vector-icons (Feather icon set)\n"
        "- react-native-chart-kit for reports\n"
        "- Same API client pattern as web, pointing to https://api.nyumbayangu.online/api\n"
        "- Role-based navigation: PublicStack, LandlordTabs, TenantTabs, AdminTabs\n"
        "- All screens mirror the web frontend functionality"
    )

    heading(doc, "Request Flow", level=3)
    doc.add_paragraph(
        "Browser/App → fetch with JWT in Authorization header\n"
        "  → chi router /api/*\n"
        "  → auth + role middleware gate\n"
        "  → Handler (one per domain)\n"
        "  → raw SQL via pgx\n"
        "  → PostgreSQL"
    )

    heading(doc, "Authentication", level=3)
    doc.add_paragraph(
        "Two independent ways to end up with a session token, both returning { token, user }:\n\n"
        "1. Phone + password — POST /auth/register, POST /auth/login. Landlords/agents who submit a "
        "NIDA number or business license at registration start in \"pending\" verification automatically.\n\n"
        "2. Email OTP — POST /auth/otp/send emails a random 6-digit code (10-minute expiry, single use, "
        "sent via Resend); POST /auth/otp/verify checks it and auto-creates an account on first use."
    )
    doc.add_page_break()

    # ---- 05 ROADMAP ----
    heading(doc, "05 — ROADMAP")
    heading(doc, "Scope & Phasing", level=2)

    table(doc,
        ["ID", "Description", "Status"],
        [
            ["Phase 1", "Core marketplace + management suite\n"
             "Guest browsing/search, listings (rent + sale + land), wishlist, inquiries, "
             "landlord property/unit/lease/payment management, arrears, reports, "
             "subscription pricing, admin KYC queue.", "SHIPPED"],
            ["Phase 2", "Marketplace polish (frontend-only)\n"
             "Price/bedroom filters + sort, recently-viewed & similar listings, "
             "seller view-count stats, TZS/USD display toggle.", "SHIPPED"],
            ["Phase 3", "Trust, payments & tenant experience\n"
             "Amenities checklist, maintenance requests, under-renovation status. "
             "Mobile app (iOS + Android).", "SHIPPED"],
            ["Phase 4", "Growth & monetization\n"
             "Real M-Pesa/aggregator integration, subscription billing, featured listings, "
             "phone validation, booking/viewing system, sign-in notifications, "
             "property photo gallery improvements, tenant ratings.", "NEEDS BACKEND"],
            ["Phase 5", "Scale & expansion\n"
             "Multi-city support (Dodoma, Arusha, Mwanza), agent management features, "
             "cloud media storage, offline-capable mobile, push notifications, "
             "tenant insurance, market data analytics.", "PLANNED"],
        ]
    )
    doc.add_page_break()

    # ---- 06 REQUIREMENTS: PUBLIC & TENANT ----
    heading(doc, "06 — REQUIREMENTS: Public & Tenant Portal")

    heading(doc, "Marketplace (no login required)", level=2)
    table(doc,
        ["ID", "Requirement", "Status"],
        [
            ["FR-1.1", "Search & filter listings\n"
             "By district, ward, property type, purpose (rent/sale), price range, "
             "bedrooms, and sort order (newest / price).", "SHIPPED"],
            ["FR-1.2", "Listing detail page\n"
             "Gallery, price, room breakdown, district/ward, verification badge, "
             "land-plot fields (acreage, title deed status), amenities where applicable.", "SHIPPED"],
            ["FR-1.3", "Contact / inquiry form\n"
             "Public, unauthenticated POST today — name, phone, message reach the "
             "landlord's inquiry inbox.", "OPEN DECISION\n(auth-gate it?)"],
            ["FR-1.4", "Save / wishlist listings\n"
             "Works pre-login via local storage, merges into the account on sign-in.", "SHIPPED"],
            ["FR-1.5", "Recently viewed & similar listings\n"
             "Local-storage history strip and a same-district/type similar rail on "
             "each detail page.", "SHIPPED"],
            ["FR-1.6", "TZS / USD display toggle\n"
             "Static approximate rate, display-only — never sent to the backend.", "SHIPPED"],
            ["FR-1.7", "Local amenities on listings\n"
             "LUKU vs. shared meter, DAWASA/borehole water, fenced compound, security "
             "gate, balcony, paved parking, master-bedroom count. Stored as columns on "
             "units table + 3 property-level fields.", "SHIPPED"],
            ["FR-1.8", "Book a viewing / inspection\n"
             "Structured date+time request, distinct from a general inquiry message.", "NEEDS BACKEND"],
            ["FR-1.9", "Listing photo gallery (multiple images)\n"
             "Swipeable carousel on detail page, landlord uploads multiple photos "
             "per listing.", "SHIPPED\n(basic)"],
            ["FR-1.10", "Map view of listings\n"
             "Show listings on a map with pin markers, filterable by district. "
             "Requires lat/lng on properties (columns exist but unused).", "NEW\nNEEDS BACKEND"],
            ["FR-1.11", "Share listing via WhatsApp / social\n"
             "Deep link to listing detail page, shareable via native share sheet on mobile "
             "and copy-link on web.", "NEW\nFRONTEND-ONLY"],
        ]
    )

    heading(doc, "Tenant Portal (login required, role: tenant)", level=2)
    table(doc,
        ["ID", "Requirement", "Status"],
        [
            ["FR-2.1", "My leases\n"
             "Active/past leases, term, rent amount, digital-signature status.", "SHIPPED"],
            ["FR-2.2", "Rent & payment history\n"
             "Payment schedule per lease with status (pending / paid / partial / overdue).", "SHIPPED"],
            ["FR-2.3", "Self-service profile\n"
             "Tenant edits their own contact/NIDA/emergency-contact details.", "SHIPPED"],
            ["FR-2.4", "Sign a lease digitally\n"
             "Tenant-side signature capture (name, timestamp, IP) against a "
             "landlord-issued lease.", "SHIPPED"],
            ["FR-2.5", "Maintenance request form\n"
             "Tenant submits title, description, priority (low/medium/high/urgent). "
             "Request linked to their leased unit. Shows list of all submitted requests "
             "with status badges.", "SHIPPED"],
            ["FR-2.6", "Maintenance request notifications\n"
             "Tenant gets notified when landlord updates request status "
             "(open > in_progress > completed > closed).", "NEW\nNEEDS BACKEND"],
            ["FR-2.7", "Rent payment reminders\n"
             "Push notification / SMS when rent due date is approaching (3 days before, "
             "day of, 3 days overdue).", "NEW\nNEEDS BACKEND"],
            ["FR-2.8", "Lease renewal request\n"
             "Tenant can request lease renewal before expiry. Landlord gets notified.", "NEW\nNEEDS BACKEND"],
            ["FR-2.9", "Payment receipt download\n"
             "Tenant can download a PDF receipt for each completed payment.", "NEW\nNEEDS BACKEND"],
        ]
    )
    doc.add_page_break()

    # ---- 07 REQUIREMENTS: LANDLORD ----
    heading(doc, "07 — REQUIREMENTS: Landlord/Agent Portal")
    table(doc,
        ["ID", "Requirement", "Status"],
        [
            ["FR-3.1", "Add property & units\n"
             "Property (building/house/plot) then per-unit records with label, beds/baths, "
             "rent, and amenities (meter type, water source, fence, gate, balcony, parking, "
             "master bedrooms).", "SHIPPED"],
            ["FR-3.2", "Publish a listing\n"
             "From a vacant unit or a land property — goes live immediately.", "SHIPPED"],
            ["FR-3.3", "Image uploads\n"
             "Multipart upload, JPEG/PNG/WebP, 5 MB cap, stored on API server disk.", "SHIPPED"],
            ["FR-3.4", "Occupancy status\n"
             "Unit status: vacant / occupied / maintenance / under_renovation. "
             "Flips to occupied automatically when a lease is created. Under-renovation "
             "status shown with distinct badge.", "SHIPPED"],
            ["FR-3.5", "Tenant & lease tracking\n"
             "Look up a tenant by phone, draft a lease against a unit, track start/end "
             "dates and status. Lease in Swahili or English.", "SHIPPED"],
            ["FR-3.6", "Record a rent payment\n"
             "Manual (cash/bank) entry, or a simulated M-Pesa STK push. "
             "M-Pesa is mocked, not live.", "SHIPPED"],
            ["FR-3.7", "Arrears tracking\n"
             "Overdue schedules surfaced per landlord, auto-flagged past due date.", "SHIPPED"],
            ["FR-3.8", "Portfolio reports\n"
             "Occupancy summary and per-building income/arrears breakdown with charts.", "SHIPPED"],
            ["FR-3.9", "Subscription view\n"
             "Read-only: current tier, unit count, computed monthly price.", "SHIPPED"],
            ["FR-3.10", "Listing view-count stats\n"
             "Landlord sees how many times each published listing was viewed.", "SHIPPED"],
            ["FR-3.11", "Maintenance request inbox\n"
             "Landlord sees all requests from tenants across their properties. "
             "Can update status (open > in_progress > completed > closed). "
             "Shows tenant name, unit, priority badge.", "SHIPPED"],
            ["FR-3.12", "Real M-Pesa / aggregator integration\n"
             "Swap the mocked STK-push handler for Selcom, ClickPesa, or AzamPay. "
             "Request/response contract is already shaped to match.", "NEEDS BACKEND"],
            ["FR-3.13", "Manual unit status toggle\n"
             "Landlord can manually change unit status (vacant/maintenance/under_renovation) "
             "without waiting for lease events.", "NEW\nNEEDS BACKEND"],
            ["FR-3.14", "Tenant communication log\n"
             "In-app messaging between landlord and tenant, replacing WhatsApp for "
             "property-related communication.", "NEW\nNEEDS BACKEND"],
            ["FR-3.15", "Lease expiry alerts\n"
             "Notify landlord 30/14/7 days before a lease expires so they can "
             "renew or plan for vacancy.", "NEW\nNEEDS BACKEND"],
            ["FR-3.16", "Bulk payment import\n"
             "Upload a CSV of M-Pesa transaction records to auto-match payments "
             "to tenants and schedules.", "NEW\nNEEDS BACKEND"],
            ["FR-3.17", "Tenant rating / review\n"
             "Landlord can privately rate a tenant (payment reliability, property care) "
             "visible only to landlords — not public.", "NEW\nNEEDS BACKEND"],
            ["FR-3.18", "Property documents vault\n"
             "Upload and store property-related documents (title deeds, permits, "
             "contracts) linked to each property.", "NEW\nNEEDS BACKEND"],
        ]
    )
    doc.add_page_break()

    # ---- 08 REQUIREMENTS: ADMIN ----
    heading(doc, "08 — REQUIREMENTS: Admin Portal")
    table(doc,
        ["ID", "Requirement", "Status"],
        [
            ["FR-4.1", "Landlord/agent KYC queue\n"
             "Review NIDA number or business license, approve or reject.", "SHIPPED"],
            ["FR-4.2", "Property verification queue\n"
             "Approve or reject a submitted property; earns the \"Verified\" badge.", "SHIPPED"],
            ["FR-4.3", "Listing moderation queue\n"
             "Doesn't exist separately — a listing goes live the moment a landlord "
             "publishes it. Decide whether property/user verification is enough "
             "gatekeeping, or a listing needs its own review step.", "OPEN DECISION"],
            ["FR-4.4", "Platform analytics dashboard\n"
             "Admin sees total users, listings, leases, revenue across the platform. "
             "Growth metrics over time.", "NEW\nNEEDS BACKEND"],
            ["FR-4.5", "User management\n"
             "Admin can view all users, suspend/unsuspend accounts, reset passwords.", "NEW\nNEEDS BACKEND"],
            ["FR-4.6", "Reported listings / disputes\n"
             "Users can flag suspicious listings. Admin reviews and takes action.", "NEW\nNEEDS BACKEND"],
        ]
    )
    doc.add_page_break()

    # ---- 09 REQUIREMENTS: MOBILE ----
    heading(doc, "09 — REQUIREMENTS: Mobile App (NEW)")
    doc.add_paragraph(
        "A React Native + Expo mobile app has been built mirroring all web frontend functionality. "
        "Both iOS and Android are supported from a single codebase."
    )
    table(doc,
        ["ID", "Requirement", "Status"],
        [
            ["FR-5.1", "Public marketplace browsing\n"
             "Home screen, listing search/filter, listing detail, wishlist — "
             "all available without login.", "SHIPPED"],
            ["FR-5.2", "User authentication\n"
             "Login (phone+password), Register (with OTP), role-based navigation.", "SHIPPED"],
            ["FR-5.3", "Landlord dashboard\n"
             "Overview with stats, properties & units with amenities, tenants, "
             "leases, lease detail, reports (charts), subscription, maintenance inbox.", "SHIPPED"],
            ["FR-5.4", "Tenant portal\n"
             "Lease view, e-signature, M-Pesa payment, payment schedule/history, "
             "maintenance request form.", "SHIPPED"],
            ["FR-5.5", "Admin verification\n"
             "Pending users and properties, approve/reject.", "SHIPPED"],
            ["FR-5.6", "Push notifications\n"
             "Rent reminders, maintenance updates, lease expiry alerts — "
             "via Expo Notifications.", "NEW\nNEEDS BACKEND"],
            ["FR-5.7", "Offline support\n"
             "Cache recent data for viewing when offline. Queue actions "
             "(payment log, maintenance request) for sync when back online.", "NEW\nNEEDS BACKEND"],
            ["FR-5.8", "Camera integration for uploads\n"
             "Landlord takes property photos directly from the app camera.", "NEW\nFRONTEND-ONLY"],
            ["FR-5.9", "Biometric login\n"
             "Face ID / fingerprint unlock after initial login.", "NEW\nFRONTEND-ONLY"],
        ]
    )
    doc.add_page_break()

    # ---- 10 DATA MODEL ----
    heading(doc, "10 — DATA MODEL")
    heading(doc, "Contracts", level=2)
    doc.add_paragraph(
        "These are the real shapes the Go API returns today — not a proposal. Field names below are "
        "what the frontend TypeScript types and the Postgres schema both agree on."
    )

    heading(doc, "Property", level=3)
    doc.add_paragraph(
        "{\n"
        '  "id": "a1b2...",\n'
        '  "owner_id": "u9f8...",\n'
        '  "title": "Modern 2-Bedroom Apartments",\n'
        '  "property_type": "apartment",  // apartment | house | room | office | shop | land\n'
        '  "district": "Kinondoni",\n'
        '  "ward": "Mikocheni B",\n'
        '  "city": "Dar es Salaam",\n'
        '  "land_size_acres": null,\n'
        '  "title_deed_status": null,\n'
        '  "has_fence": false,\n'
        '  "has_security_gate": true,\n'
        '  "has_parking": true,\n'
        '  "verification": "verified",  // unverified | pending | verified | rejected\n'
        '  "created_at": "2026-08-20T10:00:00Z"\n'
        "}"
    )

    heading(doc, "Unit", level=3)
    doc.add_paragraph(
        "{\n"
        '  "id": "unit_441",\n'
        '  "property_id": "a1b2...",\n'
        '  "unit_label": "Flat 2B",\n'
        '  "bedrooms": 2, "bathrooms": 1, "size_sqm": null,\n'
        '  "rent_amount": 700000,\n'
        '  "status": "vacant",  // vacant | occupied | maintenance | under_renovation\n'
        '  "meter_type": "LUKU",  // LUKU | shared | prepaid | null\n'
        '  "water_source": "DAWASA",  // DAWASA | borehole | well | null\n'
        '  "has_fence": true,\n'
        '  "has_security_gate": true,\n'
        '  "has_balcony": false,\n'
        '  "master_bedrooms": 1,\n'
        '  "has_parking": true,\n'
        '  "created_at": "2026-08-20T10:00:00Z"\n'
        "}"
    )

    heading(doc, "Listing (public, joined view)", level=3)
    doc.add_paragraph(
        "{\n"
        '  "id": "l_9012",\n'
        '  "property_id": "a1b2...", "unit_id": "unit_441",\n'
        '  "purpose": "rent",  // rent | sale\n'
        '  "title": "Flat 2B - TSh 700,000/month",\n'
        '  "price": 700000, "price_period": "month",\n'
        '  "status": "active",  // active | paused | let | sold | expired\n'
        '  "view_count": 47,\n'
        '  "district": "Kinondoni", "ward": "Mikocheni B",\n'
        '  "bedrooms": 2, "bathrooms": 1,\n'
        '  "images": ["/uploads/8f2c1a...jpg"]\n'
        "}"
    )

    heading(doc, "Lease", level=3)
    doc.add_paragraph(
        "{\n"
        '  "id": "lease_9901",\n'
        '  "unit_id": "unit_441",\n'
        '  "tenant_id": "t_781",\n'
        '  "landlord_id": "u9f8...",\n'
        '  "language": "sw",  // sw | en\n'
        '  "start_date": "2026-09-01", "end_date": "2027-02-28",\n'
        '  "rent_amount": 700000,\n'
        '  "advance_months": 6,\n'
        '  "deposit_amount": 700000,\n'
        '  "status": "active",  // draft | active | terminated | expired\n'
        '  "tenant_signed_at": "2026-08-30T14:02:00Z",\n'
        '  "landlord_signed_at": "2026-08-29T09:11:00Z"\n'
        "}"
    )

    heading(doc, "Maintenance Request (NEW)", level=3)
    doc.add_paragraph(
        "{\n"
        '  "id": "mr_1234",\n'
        '  "unit_id": "unit_441",\n'
        '  "tenant_user_id": "u_tenant_99",\n'
        '  "title": "Broken pipe in bathroom",\n'
        '  "description": "Water leaking from under the sink",\n'
        '  "priority": "high",  // low | medium | high | urgent\n'
        '  "status": "open",  // open | in_progress | completed | closed\n'
        '  "unit_label": "Flat 2B",\n'
        '  "property_title": "Modern 2-Bedroom Apartments",\n'
        '  "tenant_name": "John Mwalimu",\n'
        '  "created_at": "2026-08-23T10:00:00Z",\n'
        '  "updated_at": "2026-08-23T10:00:00Z"\n'
        "}"
    )
    doc.add_page_break()

    # ---- 11 API REFERENCE ----
    heading(doc, "11 — API REFERENCE")
    doc.add_paragraph(
        "Grouped by who can call it. Everything under /api; auth is a bearer JWT unless marked public."
    )

    heading(doc, "Public — no auth", level=2)
    table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["POST", "/auth/register", "Phone + password sign-up"],
            ["POST", "/auth/login", "Phone + password sign-in"],
            ["POST", "/auth/otp/send", "Email a 6-digit login code"],
            ["POST", "/auth/otp/verify", "Redeem the code, auto-creates account"],
            ["GET", "/listings", "Browse/search/filter"],
            ["GET", "/listings/{id}", "Detail + increments view_count"],
            ["POST", "/listings/{id}/inquiries", "Contact-lister form"],
            ["GET", "/properties/{id}", "Single property (public view)"],
        ]
    )

    heading(doc, "Any authenticated user", level=2)
    table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["GET", "/me", "Current user profile"],
            ["GET", "/favorites", "Saved listings (full objects)"],
            ["GET", "/favorites/ids", "Saved listing IDs only (cheap check)"],
            ["POST", "/favorites", "Save a listing"],
            ["POST", "/favorites/sync", "Merge local-storage saves into account"],
            ["DELETE", "/favorites/{listingID}", "Unsave a listing"],
            ["GET", "/leases/{id}/schedules", "Payment schedule for a lease"],
            ["GET", "/leases/{id}/payments", "Payment history for a lease"],
            ["GET", "/leases/{id}/document", "Download lease document"],
            ["POST", "/leases/{id}/sign", "Digital signature"],
            ["POST", "/payments/mpesa/stk-push", "M-Pesa STK push (simulated)"],
        ]
    )

    heading(doc, "Landlord / Agent", level=2)
    table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["POST", "/properties", "Create property (with amenities)"],
            ["GET", "/properties", "List my properties"],
            ["POST", "/properties/{id}/units", "Create unit (with amenities)"],
            ["GET", "/properties/{id}/units", "List units for property"],
            ["GET", "/units", "All units across my properties"],
            ["POST", "/listings", "Publish a listing"],
            ["POST", "/uploads", "Image upload (multipart)"],
            ["GET", "/inquiries", "Inquiry inbox"],
            ["GET", "/tenants", "My tenants"],
            ["GET", "/tenants/lookup", "Find tenant by phone"],
            ["POST", "/leases", "Draft a lease"],
            ["GET", "/leases", "List my leases"],
            ["POST", "/payments/manual", "Log cash/bank payment"],
            ["GET", "/arrears", "Overdue schedules"],
            ["GET", "/reports/summary", "Portfolio KPIs"],
            ["GET", "/reports/by-building", "Income per building"],
            ["GET", "/subscription", "Current tier & price"],
            ["GET", "/maintenance-requests", "Tenant maintenance requests"],
            ["PATCH", "/maintenance-requests/{id}", "Update request status"],
        ]
    )

    heading(doc, "Tenant", level=2)
    table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["GET", "/tenant/me", "My tenant profile"],
            ["POST", "/tenant/profile", "Create/update my profile"],
            ["GET", "/tenant/leases", "My leases"],
            ["POST", "/tenant/maintenance-requests", "Submit maintenance request"],
            ["GET", "/tenant/maintenance-requests", "My maintenance requests"],
        ]
    )

    heading(doc, "Admin", level=2)
    table(doc,
        ["Method", "Path", "Purpose"],
        [
            ["GET", "/admin/pending-users", "KYC queue"],
            ["GET", "/admin/pending-properties", "Property verification queue"],
            ["POST", "/admin/users/{id}/decide", "Approve/reject a user"],
            ["POST", "/admin/properties/{id}/decide", "Approve/reject a property"],
        ]
    )
    doc.add_page_break()

    # ---- 12 NON-FUNCTIONAL REQUIREMENTS ----
    heading(doc, "12 — NON-FUNCTIONAL REQUIREMENTS")
    table(doc,
        ["ID", "Requirement", "Status"],
        [
            ["NFR-1", "Localization\n"
             "TZS is the source of truth everywhere money is stored. Lease documents "
             "render in Swahili or English per lease. USD is display-only.", "SHIPPED"],
            ["NFR-2", "Geographic scope\n"
             "Dar es Salaam only — 5 districts, ward-level granularity. Expanding "
             "beyond Dar is a schema decision.", "OPEN DECISION"],
            ["NFR-3", "Auth security\n"
             "Bcrypt password hashes, JWT with 7-day expiry, role-gated route groups "
             "server-side (not just hidden UI).", "SHIPPED"],
            ["NFR-4", "Upload validation\n"
             "Server-side content-type sniffing, 5 MB cap, JPEG/PNG/WebP only.", "SHIPPED"],
            ["NFR-5", "Phone number validation\n"
             "Nothing enforces a Tanzanian format (+255 / 06 / 07) anywhere today.", "NEEDS BACKEND"],
            ["NFR-6", "Media storage & scaling\n"
             "Uploaded images sit on the API server's local disk. Fine for one VPS; "
             "becomes a constraint with multiple instances.", "OPEN DECISION"],
            ["NFR-7", "Mobile performance\n"
             "App should load within 3 seconds on 3G network. Images lazy-loaded. "
             "Lists use FlatList for virtualization.", "SHIPPED"],
            ["NFR-8", "Data backup\n"
             "Daily PostgreSQL backup to external storage. No backup system exists yet.", "NEW\nNEEDS BACKEND"],
            ["NFR-9", "Rate limiting\n"
             "API rate limiting to prevent abuse (especially on public endpoints "
             "like OTP send, listings browse).", "NEW\nNEEDS BACKEND"],
            ["NFR-10", "HTTPS everywhere\n"
             "All API and web traffic over TLS via nginx + Let's Encrypt.", "SHIPPED"],
        ]
    )
    doc.add_page_break()

    # ---- 13 CORRECTIONS ----
    heading(doc, "13 — CORRECTIONS TO ORIGINAL DRAFT")
    doc.add_paragraph(
        "The architecture note originally drafted described a system close in spirit to what's built, "
        "but several specifics didn't match the real codebase. These corrections were documented in "
        "SRS v1.0 and remain valid:"
    )
    table(doc,
        ["What was written", "What is actually built"],
        [
            ["JavaScript (.jsx)", "TypeScript (.tsx) throughout"],
            ["Axios with interceptors", "Hand-written fetch wrapper in lib/api.ts"],
            ["Context API / Zustand", "Context API only, one provider per concern"],
            ["region / neighborhood_mtaa", "district and ward (Dar es Salaam only)"],
            ["listing_type field", "purpose field (\"rent\" | \"sale\") on Listing"],
            ["unit_number_name, master_bedrooms, has_luku", "unit_label + amenity fields now exist (v2.0)"],
            ["Enum string \"6_months\"", "advance_months as plain integer"],
            ["Tenant embedded in lease", "Separate tenants table, linked by tenant_id"],
            ["Cloudinary/S3 presigned URLs", "Local disk uploads on VPS"],
            ["Suffix style \"650,000 TZS\"", "Intl.NumberFormat (\"TSh 650,000\")"],
            ["Phone validation exists", "Not implemented yet (see NFR-5)"],
        ]
    )
    doc.add_page_break()

    # ---- 14 OPEN DECISIONS & BACKLOG ----
    heading(doc, "14 — OPEN DECISIONS & BACKLOG")
    doc.add_paragraph(
        "Everything here needs either new backend surface, a joint decision, or both — nothing in "
        "this list should get built solo without the other person weighing in first."
    )
    table(doc,
        ["ID", "Decision", "Status"],
        [
            ["DEC-1", "Auth wall on contact / booking\n"
             "Should \"Contact lister\" require login? Currently public.", "OPEN DECISION"],
            ["DEC-2", "Amenities schema\n"
             "New columns on Property/Unit for LUKU, DAWASA, fence, gate, balcony, "
             "parking, master bedrooms.", "SHIPPED (v2.0)"],
            ["DEC-3", "Maintenance requests\n"
             "New table + handler + tenant form + landlord inbox.", "SHIPPED (v2.0)"],
            ["DEC-4", "Real M-Pesa integration\n"
             "Replace mocked STK-push with Selcom, ClickPesa, or AzamPay. "
             "Pick a provider before building.", "NEEDS BACKEND"],
            ["DEC-5", "Subscription billing\n"
             "Pricing computes correctly, but nothing charges a landlord yet.", "NEEDS BACKEND"],
            ["DEC-6", "Listing moderation\n"
             "Is property/user KYC sufficient, or do listings need their own "
             "admin review step?", "OPEN DECISION"],
            ["DEC-7", "Phone number validation\n"
             "Enforce Tanzanian formats (+255 / 06 / 07) server-side.", "NEEDS BACKEND"],
            ["DEC-8", "Sign-in notification email\n"
             "\"Someone signed into your account\" alert.", "NEEDS BACKEND"],
            ["DEC-9", "Media storage strategy\n"
             "Move off local-disk to S3/Cloudinary before scaling.", "OPEN DECISION"],
            ["DEC-10", "Multi-city expansion\n"
             "Adding Dodoma, Arusha, Mwanza requires a region/city schema change.", "OPEN DECISION"],
            ["DEC-11", "Push notifications infrastructure\n"
             "Choose provider: Expo Push, Firebase FCM, or OneSignal.", "OPEN DECISION"],
            ["DEC-12", "In-app messaging vs. WhatsApp\n"
             "Build tenant-landlord chat, or keep using WhatsApp links?", "OPEN DECISION"],
            ["DEC-13", "Tenant insurance partnerships\n"
             "Partner with an insurer for tenant deposit protection.", "OPEN DECISION"],
            ["DEC-14", "Data analytics / market reports\n"
             "Sell anonymized rental market data to developers/investors?", "OPEN DECISION"],
        ]
    )

    # Save
    path = os.path.join(OUT_DIR, "Nyumba_Yangu_SRS_v2.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_srs()
