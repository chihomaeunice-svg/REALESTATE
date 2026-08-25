#!/usr/bin/env python3
"""Generate all Nyumba Yangu business documents as Word (.docx) files."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

BRAND_COLOR = RGBColor(0x1D, 0x4E, 0x89)


def set_style(doc):
    """Apply consistent styling."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_title_page(doc, title, subtitle=""):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_COLOR
    run.bold = True

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Nyumba Yangu Ltd.")
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Dar es Salaam, Tanzania\nwww.nyumbayangu.online\n\nAugust 2026")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Prepared by:\nEunice Chihoma — Co-Founder & CEO\nErick Mkingule — Co-Founder & CTO")
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_COLOR


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph("")
    return table


# ============================================================
# 1. BUSINESS PLAN
# ============================================================
def create_business_plan():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "BUSINESS PLAN", "Tanzania's Hybrid Real Estate Platform")

    # Table of Contents
    add_heading(doc, "Table of Contents")
    toc_items = [
        "1. Executive Summary",
        "2. Company Overview",
        "3. Problem Statement",
        "4. Solution",
        "5. Market Analysis",
        "6. Business Model & Revenue",
        "7. Product & Technology",
        "8. Go-To-Market Strategy",
        "9. Operations Plan",
        "10. Team",
        "11. Financial Projections",
        "12. Risk Analysis",
        "13. Milestones & Roadmap",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "Nyumba Yangu (\"My House\" in Swahili) is a technology platform that solves the broken "
        "rental market in Tanzania. Today, landlords manage properties using paper ledgers and WhatsApp groups, "
        "tenants have no visibility into payments or lease terms, and property seekers rely on unreliable "
        "brokers (dalali) who charge high fees for poor service."
    )
    doc.add_paragraph(
        "Nyumba Yangu combines three products in one platform:\n"
        "- A public property marketplace where seekers find verified listings\n"
        "- A landlord management suite for properties, tenants, leases, and payments\n"
        "- A tenant portal for lease signing, rent payment via M-Pesa, and maintenance requests"
    )
    doc.add_paragraph(
        "The platform is live at nyumbayangu.online with both web and mobile (iOS/Android) apps, "
        "serving the Dar es Salaam market. Revenue comes from tiered subscriptions for landlords "
        "who manage multiple properties."
    )
    doc.add_paragraph(
        "Tanzania has over 1.5 million rental units in urban areas, with Dar es Salaam alone accounting "
        "for roughly 500,000 units. The formal property management market is virtually untouched by technology. "
        "Nyumba Yangu aims to capture 5% of the Dar es Salaam market within 3 years, generating annual "
        "recurring revenue of TZS 1.2 billion (approx. USD 460,000)."
    )

    # 2. Company Overview
    add_heading(doc, "2. Company Overview")
    add_table(doc,
        ["Item", "Detail"],
        [
            ["Company Name", "Nyumba Yangu Ltd."],
            ["Location", "Dar es Salaam, Tanzania"],
            ["Founded", "2026"],
            ["Industry", "PropTech / Real Estate Technology"],
            ["Website", "https://nyumbayangu.online"],
            ["Legal Structure", "Private Limited Company (Tanzania)"],
            ["Stage", "Live product, pre-revenue"],
        ]
    )
    doc.add_paragraph(
        "Nyumba Yangu was co-founded by Erick Mkingule (CTO) and Eunice Chihoma (CEO) "
        "to modernize property management in Tanzania. The team combines deep software engineering "
        "expertise with hands-on knowledge of the local real estate market."
    )

    # 3. Problem Statement
    add_heading(doc, "3. Problem Statement")
    add_heading(doc, "For Property Seekers", level=2)
    doc.add_paragraph(
        "Finding a rental in Dar es Salaam is frustrating. Most listings are posted on WhatsApp groups, "
        "Instagram, or through brokers (dalali) who charge 1-3 months' rent as fees. There is no centralized, "
        "trusted platform with verified listings, transparent pricing, and direct landlord contact."
    )
    add_heading(doc, "For Landlords", level=2)
    doc.add_paragraph(
        "Landlords managing 5-50+ units track everything manually - paper receipts, Excel spreadsheets, "
        "or memory. Common problems include:\n"
        "- No systematic rent collection tracking\n"
        "- Difficulty identifying which tenants are in arrears\n"
        "- No digital lease records\n"
        "- No way to receive maintenance requests from tenants\n"
        "- Time wasted on phone calls and physical visits for routine management"
    )
    add_heading(doc, "For Tenants", level=2)
    doc.add_paragraph(
        "Tenants have no digital record of their lease terms or payment history. "
        "Disputes are common because everything is verbal or on paper. "
        "Paying rent requires physical cash handover or informal M-Pesa transfers "
        "with no receipt linked to a specific lease period."
    )

    # 4. Solution
    add_heading(doc, "4. Solution")
    doc.add_paragraph("Nyumba Yangu is a three-layer platform:")

    add_heading(doc, "Layer 1: Public Marketplace", level=2)
    doc.add_paragraph(
        "A searchable directory of verified property listings. Seekers can filter by district, "
        "ward, property type, price range, and bedrooms. Each listing shows photos, amenities "
        "(electricity type, water source, parking, security), and direct landlord contact. "
        "No broker fees."
    )

    add_heading(doc, "Layer 2: Landlord Management Suite", level=2)
    doc.add_paragraph(
        "A complete dashboard for landlords to:\n"
        "- Register properties and individual units with full amenity details\n"
        "- Publish listings to the public marketplace in one click\n"
        "- Create and manage tenant profiles\n"
        "- Generate digital leases (Swahili or English) with e-signature\n"
        "- Track rent payments with automatic schedule generation\n"
        "- View income reports, arrears, and building-level analytics\n"
        "- Receive and manage maintenance requests from tenants"
    )

    add_heading(doc, "Layer 3: Tenant Portal", level=2)
    doc.add_paragraph(
        "Tenants get their own login to:\n"
        "- View their lease terms and download the lease document\n"
        "- Sign leases electronically\n"
        "- Pay rent via M-Pesa STK push (money is deducted from their phone)\n"
        "- View payment history and upcoming schedule\n"
        "- Submit maintenance requests (broken pipe, faulty lock, etc.)"
    )

    # 5. Market Analysis
    add_heading(doc, "5. Market Analysis")
    add_heading(doc, "Market Size", level=2)
    add_table(doc,
        ["Metric", "Value", "Source"],
        [
            ["Tanzania urban population", "~20 million", "NBS 2022 Census"],
            ["Dar es Salaam population", "~6.4 million", "NBS 2022 Census"],
            ["Estimated urban rental units (Tanzania)", "1.5 million+", "NHC / World Bank estimates"],
            ["Estimated rental units (Dar es Salaam)", "~500,000", "Industry estimate"],
            ["Average rent (Dar es Salaam, 1-2BR)", "TZS 150,000 - 500,000/month", "Market research"],
            ["Annual rental market value (DSM)", "TZS 1+ trillion", "Calculated"],
        ]
    )

    add_heading(doc, "Target Market", level=2)
    doc.add_paragraph(
        "Primary: Landlords in Dar es Salaam managing 3-50 rental units (apartments, houses, rooms). "
        "These are individuals or small businesses, not institutional players. They currently use no "
        "technology for property management.\n\n"
        "Secondary: Property seekers in Dar es Salaam looking for rentals. "
        "Young professionals (25-40) with smartphones who are frustrated with the dalali system.\n\n"
        "Tertiary: Real estate agents who manage properties on behalf of owners."
    )

    add_heading(doc, "Competitive Landscape", level=2)
    doc.add_paragraph(
        "Tanzania has very few PropTech players compared to Kenya or Nigeria:"
    )
    add_table(doc,
        ["Competitor", "Type", "Strength", "Weakness"],
        [
            ["WhatsApp / Social Media", "Informal", "Everyone uses it", "No structure, no verification, scams"],
            ["Dalali (brokers)", "Traditional", "Local knowledge", "High fees (1-3 months rent), unreliable"],
            ["Tanzania Property Centre", "Listing site", "Established brand", "No management tools, no tenant portal"],
            ["BeiYaUziki", "Listing site", "Local focus", "Listings only, no SaaS"],
            ["Nyumba Yangu", "Full platform", "Marketplace + Management + Tenant portal", "New entrant, building trust"],
        ]
    )
    doc.add_paragraph(
        "Our key differentiator: we are not just a listing site. We combine the marketplace with "
        "operational tools that make landlords' lives easier. Once a landlord uses our management suite, "
        "switching costs are high because their tenants, leases, and payment history live on our platform."
    )

    # 6. Business Model
    add_heading(doc, "6. Business Model & Revenue")
    add_heading(doc, "Revenue Model: SaaS Subscriptions", level=2)
    doc.add_paragraph(
        "Nyumba Yangu uses a freemium subscription model for landlords. "
        "The marketplace (Layer 1) is free for everyone. "
        "Landlords pay monthly subscriptions based on the number of units they manage."
    )
    add_table(doc,
        ["Tier", "Units", "Price (TZS/month)", "Price (USD/month)", "Features"],
        [
            ["Free", "1-3", "0", "$0", "Basic listings, 1 property"],
            ["Starter", "4-10", "25,000", "~$10", "Full management suite, reports"],
            ["Growth", "11-30", "60,000", "~$23", "Everything + priority support"],
            ["Scale", "31+", "120,000", "~$46", "Everything + API access, white-label"],
        ]
    )

    add_heading(doc, "Future Revenue Streams", level=2)
    doc.add_paragraph(
        "- Featured listings: Landlords pay to boost visibility (TZS 10,000-50,000 per listing)\n"
        "- Transaction fees: Small percentage on M-Pesa rent payments (1-2%)\n"
        "- Insurance partnerships: Tenant deposit insurance\n"
        "- Data services: Anonymized market data for real estate developers and investors"
    )

    add_heading(doc, "Unit Economics", level=2)
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Average Revenue Per User (ARPU)", "TZS 50,000/month (~$19)"],
            ["Customer Acquisition Cost (CAC)", "TZS 30,000 (~$12)"],
            ["Lifetime Value (LTV, 24-month avg retention)", "TZS 1,200,000 (~$460)"],
            ["LTV:CAC Ratio", "40:1"],
            ["Gross Margin", "85%+ (SaaS, minimal COGS)"],
            ["Payback Period", "< 1 month"],
        ]
    )

    # 7. Product & Technology
    add_heading(doc, "7. Product & Technology")
    add_heading(doc, "Technology Stack", level=2)
    add_table(doc,
        ["Component", "Technology", "Why"],
        [
            ["Backend API", "Go (Golang)", "Fast, reliable, low server costs"],
            ["Database", "PostgreSQL", "Robust, proven, handles complex queries"],
            ["Web Frontend", "React + TypeScript + Vite", "Modern, fast development"],
            ["Mobile App", "React Native + Expo", "Single codebase for iOS and Android"],
            ["Payments", "M-Pesa STK Push", "How Tanzanians pay - 80%+ mobile money penetration"],
            ["Hosting", "VPS (Vultr)", "Cost-effective, Nairobi data center (low latency)"],
            ["Domain", "nyumbayangu.online", "Live and operational"],
        ]
    )

    add_heading(doc, "Platform Features (Live)", level=2)
    add_table(doc,
        ["Feature", "Status", "Users"],
        [
            ["Property & unit registration with amenities", "Live", "Landlords"],
            ["Public listing marketplace with search/filter", "Live", "Everyone"],
            ["User registration with OTP verification", "Live", "Everyone"],
            ["Digital lease creation (Swahili/English)", "Live", "Landlords"],
            ["Electronic lease signing", "Live", "Tenants"],
            ["M-Pesa rent payment (STK Push)", "Live", "Tenants"],
            ["Payment schedule & tracking", "Live", "Landlords + Tenants"],
            ["Income reports & arrears tracking", "Live", "Landlords"],
            ["Maintenance request system", "Live", "Tenants + Landlords"],
            ["Wishlist / saved listings", "Live", "Everyone"],
            ["Admin verification queue", "Live", "Admin"],
            ["Mobile app (iOS + Android)", "Live", "Everyone"],
        ]
    )

    # 8. Go-To-Market Strategy
    add_heading(doc, "8. Go-To-Market Strategy")
    add_heading(doc, "Phase 1: Dar es Salaam (Months 1-12)", level=2)
    doc.add_paragraph(
        "Focus: Onboard 200 landlords managing 2,000+ units in key districts.\n\n"
        "Channels:\n"
        "- Direct outreach: Visit apartment buildings in Kinondoni, Ilala, Ubungo, Kigamboni. "
        "Talk to landlords and caretakers in person.\n"
        "- WhatsApp marketing: Join existing property groups, share verified listings.\n"
        "- University partnerships: Partner with UDSM, Ardhi University for student housing listings.\n"
        "- Referral program: Existing landlords get 1 month free for every landlord they refer.\n"
        "- Social media: Instagram and TikTok content showing the platform in action."
    )

    add_heading(doc, "Phase 2: Expansion (Months 12-24)", level=2)
    doc.add_paragraph(
        "- Expand to Dodoma (new capital, growing rental demand)\n"
        "- Expand to Arusha and Mwanza (secondary cities)\n"
        "- Launch agent/property manager features\n"
        "- Partner with real estate companies for bulk onboarding"
    )

    add_heading(doc, "Phase 3: Scale (Months 24-36)", level=2)
    doc.add_paragraph(
        "- National coverage across all major Tanzanian cities\n"
        "- Launch premium features (featured listings, transaction fees)\n"
        "- Explore neighboring markets (Kenya, Uganda, Rwanda)"
    )

    # 9. Operations Plan
    add_heading(doc, "9. Operations Plan")
    add_heading(doc, "Team Structure (Current)", level=2)
    doc.add_paragraph(
        "- Erick Mkingule, Co-Founder & CTO: Software development, platform architecture, DevOps\n"
        "- Eunice Chihoma, Co-Founder & CEO: Business development, landlord onboarding, operations, product direction"
    )

    add_heading(doc, "Planned Hires (Year 1)", level=2)
    add_table(doc,
        ["Role", "When", "Monthly Cost (TZS)"],
        [
            ["Sales / Field Agent (x2)", "Month 3", "600,000 each"],
            ["Customer Support", "Month 6", "500,000"],
            ["Digital Marketing", "Month 6", "700,000"],
        ]
    )

    add_heading(doc, "Monthly Operating Costs", level=2)
    add_table(doc,
        ["Item", "Monthly (TZS)", "Monthly (USD)"],
        [
            ["Server hosting (VPS)", "50,000", "$19"],
            ["Domain & SSL", "5,000", "$2"],
            ["M-Pesa API fees", "Variable", "Per transaction"],
            ["SMS (OTP verification)", "20,000", "$8"],
            ["Internet & office", "100,000", "$38"],
            ["Total (before salaries)", "175,000", "~$67"],
        ]
    )
    doc.add_paragraph(
        "Note: Operating costs are extremely low because we built the technology in-house "
        "and use cost-effective hosting. This is a major advantage over competitors who rely on "
        "expensive third-party platforms."
    )

    # 10. Team
    add_heading(doc, "10. Team")
    doc.add_paragraph(
        "The founding team brings complementary skills:\n\n"
        "Erick Mkingule — Co-Founder & CTO\n"
        "Technical lead responsible for building and maintaining the full platform "
        "(Go backend, React web frontend, React Native mobile app, PostgreSQL database). "
        "Erick designed and built the entire technology stack from scratch, including the API, "
        "payment integration, and deployment infrastructure.\n\n"
        "Eunice Chihoma — Co-Founder & CEO\n"
        "Business lead responsible for operations, market research, landlord onboarding, "
        "and go-to-market execution. Eunice brings deep understanding of the Tanzanian real estate "
        "market, customer relationships, and business strategy. She drives product direction "
        "and ensures the platform solves real problems for real users."
    )

    # 11. Financial Projections
    add_heading(doc, "11. Financial Projections (3-Year)")
    add_table(doc,
        ["Metric", "Year 1", "Year 2", "Year 3"],
        [
            ["Paying landlords", "80", "350", "1,000"],
            ["Total units managed", "800", "5,000", "20,000"],
            ["Monthly revenue (TZS)", "4M", "18M", "60M"],
            ["Annual revenue (TZS)", "30M", "180M", "720M"],
            ["Annual revenue (USD)", "$11,500", "$69,000", "$277,000"],
            ["Operating costs (TZS/year)", "15M", "48M", "120M"],
            ["Net profit (TZS/year)", "15M", "132M", "600M"],
            ["Net margin", "50%", "73%", "83%"],
        ]
    )
    doc.add_paragraph(
        "Assumptions:\n"
        "- Average 10 units per paying landlord\n"
        "- Average subscription: TZS 50,000/month\n"
        "- 5% monthly churn rate\n"
        "- Year 1 focused on growth over profit\n"
        "- Costs increase with team expansion in Year 2-3"
    )

    # 12. Risk Analysis
    add_heading(doc, "12. Risk Analysis")
    add_table(doc,
        ["Risk", "Impact", "Likelihood", "Mitigation"],
        [
            ["Low landlord adoption", "High", "Medium", "Free tier + in-person onboarding + referral program"],
            ["Competitor enters market", "Medium", "Medium", "First-mover advantage, switching costs from data lock-in"],
            ["M-Pesa API issues", "Medium", "Low", "Support cash/bank recording as fallback"],
            ["Internet connectivity", "Medium", "Medium", "Mobile-first design, offline-capable features"],
            ["Regulatory changes", "Low", "Low", "Comply with Tanzania Data Protection Act, NHC guidelines"],
            ["Team capacity", "High", "Medium", "Prioritize hiring, consider technical co-founder network"],
        ]
    )

    # 13. Milestones
    add_heading(doc, "13. Milestones & Roadmap")
    add_table(doc,
        ["Timeline", "Milestone", "Status"],
        [
            ["Q3 2026", "Launch web platform (nyumbayangu.online)", "Done"],
            ["Q3 2026", "Launch mobile app (iOS + Android)", "Done"],
            ["Q3 2026", "Amenities, maintenance requests, renovation status", "Done"],
            ["Q4 2026", "Onboard first 20 landlords in Dar es Salaam", "Planned"],
            ["Q4 2026", "Launch referral program", "Planned"],
            ["Q1 2027", "100 paying landlords milestone", "Planned"],
            ["Q2 2027", "Launch featured listings (premium)", "Planned"],
            ["Q3 2027", "Expand to Dodoma", "Planned"],
            ["Q4 2027", "Launch M-Pesa transaction fees", "Planned"],
            ["2028", "National expansion, 1,000 landlords", "Planned"],
        ]
    )

    doc.add_page_break()
    add_heading(doc, "Contact")
    doc.add_paragraph(
        "Nyumba Yangu Ltd.\n"
        "Dar es Salaam, Tanzania\n"
        "Website: https://nyumbayangu.online\n"
        "Email: info@nyumbayangu.online"
    )

    path = os.path.join(OUT_DIR, "Nyumba_Yangu_Business_Plan.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# 2. PITCH DECK SCRIPT
# ============================================================
def create_pitch_deck():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "PITCH DECK", "Investor Presentation Script")

    slides = [
        ("Slide 1: Title",
         "Nyumba Yangu - Tanzania's Property Management Platform\n\n"
         "\"Imagine trying to manage 20 rental apartments with just a notebook and WhatsApp. "
         "That's the reality for 95% of landlords in Tanzania. We're changing that.\""),

        ("Slide 2: The Problem",
         "Tanzania's rental market is broken:\n\n"
         "For seekers: No trusted platform. Dalali brokers charge 1-3 months' rent. "
         "Scam listings are everywhere.\n\n"
         "For landlords: Managing properties with paper ledgers. No idea who's in arrears "
         "until it's too late. Spending hours on phone calls.\n\n"
         "For tenants: No digital lease records. No payment receipts. No way to report "
         "maintenance issues.\n\n"
         "Result: A TZS 1+ trillion rental market running on chaos."),

        ("Slide 3: The Solution",
         "Nyumba Yangu is three products in one:\n\n"
         "1. MARKETPLACE - Verified property listings with photos, amenities, direct contact. "
         "No broker fees.\n\n"
         "2. MANAGEMENT SUITE - Landlords register properties, create leases, track payments, "
         "view reports, manage maintenance requests. Everything in one dashboard.\n\n"
         "3. TENANT PORTAL - Tenants sign leases digitally, pay rent via M-Pesa, view payment "
         "history, submit maintenance requests.\n\n"
         "Think of it as Zillow + Buildium + M-Pesa, built specifically for Tanzania."),

        ("Slide 4: Demo / Product",
         "Show the live platform:\n"
         "- nyumbayangu.online (web)\n"
         "- Mobile app on phone\n\n"
         "Key flows to demonstrate:\n"
         "1. Landlord adds a property with unit amenities (LUKU, DAWASA, parking)\n"
         "2. Publishes a listing - instantly live on marketplace\n"
         "3. Creates a lease - auto-generates payment schedule\n"
         "4. Tenant opens their portal, signs lease, pays via M-Pesa\n"
         "5. Tenant submits a maintenance request\n"
         "6. Landlord sees income reports and arrears at a glance"),

        ("Slide 5: Market Size",
         "Tanzania urban rental market:\n\n"
         "- 500,000+ rental units in Dar es Salaam alone\n"
         "- 1.5M+ rental units across urban Tanzania\n"
         "- Average rent: TZS 150,000-500,000/month\n"
         "- Annual rental market: TZS 1+ trillion\n\n"
         "Our addressable market (landlords with 3+ units in DSM): ~15,000 landlords managing ~150,000 units.\n\n"
         "At TZS 50,000/month average subscription = TZS 9 billion/year addressable revenue."),

        ("Slide 6: Business Model",
         "Freemium SaaS subscriptions:\n\n"
         "Free: 1-3 units (hooks landlords into the ecosystem)\n"
         "Starter: TZS 25,000/month (4-10 units)\n"
         "Growth: TZS 60,000/month (11-30 units)\n"
         "Scale: TZS 120,000/month (31+ units)\n\n"
         "Future: Featured listings, M-Pesa transaction fees, insurance partnerships.\n\n"
         "Key insight: Once a landlord's tenants, leases, and payment history are on our platform, "
         "switching costs are extremely high."),

        ("Slide 7: Traction",
         "- Platform fully built and live (web + mobile)\n"
         "- 44 API endpoints operational\n"
         "- M-Pesa payment integration working\n"
         "- All three layers functional (marketplace + management + tenant portal)\n"
         "- Zero external funding used to date\n\n"
         "Next milestone: Onboard first 20 landlords in Dar es Salaam (Q4 2026)"),

        ("Slide 8: Competition",
         "Direct competitors in Tanzania: Almost none.\n\n"
         "Tanzania Property Centre - listings only, no management tools\n"
         "WhatsApp groups - no structure, no verification\n"
         "Dalali (brokers) - expensive, unreliable\n\n"
         "Our moat:\n"
         "1. Only platform combining marketplace + management + tenant portal\n"
         "2. Built for Tanzania (Swahili leases, M-Pesa, LUKU/DAWASA amenities)\n"
         "3. Mobile-first (80%+ of users access via phone)\n"
         "4. Data lock-in (tenant records, payment history, leases)"),

        ("Slide 9: Financial Projections",
         "Year 1: 80 paying landlords, TZS 30M revenue\n"
         "Year 2: 350 paying landlords, TZS 180M revenue\n"
         "Year 3: 1,000 paying landlords, TZS 720M revenue\n\n"
         "Gross margins: 85%+ (pure SaaS)\n"
         "Break-even: Month 8\n\n"
         "Very capital-efficient: entire platform built with under $5,000 in costs."),

        ("Slide 10: The Team",
         "Two co-founders with complementary skills:\n\n"
         "Erick Mkingule — Co-Founder & CTO\n"
         "Built the entire platform from scratch: Go backend, React web app, React Native mobile app, "
         "PostgreSQL database, M-Pesa integration, deployment pipeline. Ships fast and owns the full "
         "technical stack.\n\n"
         "Eunice Chihoma — Co-Founder & CEO\n"
         "Deep knowledge of the Tanzanian real estate market. Drives product direction, landlord "
         "relationships, business operations, and growth strategy. The voice of the customer."),

        ("Slide 11: The Ask",
         "What we need:\n\n"
         "1. Landlord onboarding support (field agents to visit properties)\n"
         "2. Marketing budget for social media and WhatsApp campaigns\n"
         "3. Working capital for first 12 months of operations\n\n"
         "What we offer:\n"
         "- A live, working product (not a prototype)\n"
         "- First-mover advantage in a massive untapped market\n"
         "- Extremely lean cost structure\n"
         "- A market where 95% of competitors are WhatsApp groups"),

        ("Slide 12: Vision",
         "Short term: Become the #1 property platform in Dar es Salaam.\n\n"
         "Medium term: Expand nationally across Tanzania.\n\n"
         "Long term: The operating system for rental property in East Africa.\n\n"
         "\"Every landlord in Tanzania should be able to manage their properties from their phone. "
         "Every tenant should be able to pay rent, sign leases, and report issues digitally. "
         "That's Nyumba Yangu.\""),
    ]

    for title, content in slides:
        add_heading(doc, title, level=1)
        doc.add_paragraph(content)
        doc.add_paragraph("")

    path = os.path.join(OUT_DIR, "Nyumba_Yangu_Pitch_Deck_Script.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# 3. FINANCIAL PROJECTIONS
# ============================================================
def create_financial_projections():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "FINANCIAL PROJECTIONS", "3-Year Financial Model")

    add_heading(doc, "Assumptions")
    add_table(doc,
        ["Assumption", "Value"],
        [
            ["Average units per landlord", "10"],
            ["Free tier cap", "3 units"],
            ["Starter price", "TZS 25,000/month"],
            ["Growth price", "TZS 60,000/month"],
            ["Scale price", "TZS 120,000/month"],
            ["Blended ARPU", "TZS 50,000/month"],
            ["Monthly churn", "5%"],
            ["New landlords/month (Y1)", "8-10"],
            ["New landlords/month (Y2)", "25-30"],
            ["New landlords/month (Y3)", "50-60"],
            ["Exchange rate", "TZS 2,600 = USD 1"],
        ]
    )

    add_heading(doc, "Year 1 - Monthly Breakdown")
    add_table(doc,
        ["Month", "New Landlords", "Total Paying", "Revenue (TZS)", "Costs (TZS)", "Net (TZS)"],
        [
            ["1", "5", "5", "250,000", "1,200,000", "-950,000"],
            ["2", "5", "10", "500,000", "1,200,000", "-700,000"],
            ["3", "8", "17", "850,000", "1,800,000", "-950,000"],
            ["4", "8", "24", "1,200,000", "1,800,000", "-600,000"],
            ["5", "8", "31", "1,550,000", "1,800,000", "-250,000"],
            ["6", "10", "39", "1,950,000", "2,500,000", "-550,000"],
            ["7", "10", "47", "2,350,000", "2,500,000", "-150,000"],
            ["8", "10", "55", "2,750,000", "2,500,000", "250,000"],
            ["9", "10", "62", "3,100,000", "2,500,000", "600,000"],
            ["10", "10", "69", "3,450,000", "2,500,000", "950,000"],
            ["11", "10", "76", "3,800,000", "2,500,000", "1,300,000"],
            ["12", "10", "80", "4,000,000", "2,500,000", "1,500,000"],
        ]
    )

    add_heading(doc, "3-Year Revenue Summary")
    add_table(doc,
        ["", "Year 1", "Year 2", "Year 3"],
        [
            ["Paying landlords (end of year)", "80", "350", "1,000"],
            ["Units under management", "800", "5,000", "20,000"],
            ["Monthly recurring revenue", "TZS 4M", "TZS 17.5M", "TZS 50M"],
            ["Annual revenue", "TZS 30M", "TZS 180M", "TZS 720M"],
            ["Annual revenue (USD)", "$11,500", "$69,000", "$277,000"],
        ]
    )

    add_heading(doc, "Cost Structure")
    add_table(doc,
        ["Cost Category", "Year 1 (TZS)", "Year 2 (TZS)", "Year 3 (TZS)"],
        [
            ["Server & infrastructure", "600,000", "2,400,000", "6,000,000"],
            ["SMS & OTP", "240,000", "600,000", "1,200,000"],
            ["Salaries (field agents)", "7,200,000", "14,400,000", "28,800,000"],
            ["Salaries (support & marketing)", "0", "14,400,000", "36,000,000"],
            ["Marketing & advertising", "3,000,000", "12,000,000", "36,000,000"],
            ["Office & admin", "1,200,000", "2,400,000", "6,000,000"],
            ["Miscellaneous", "2,760,000", "1,800,000", "6,000,000"],
            ["TOTAL COSTS", "15,000,000", "48,000,000", "120,000,000"],
        ]
    )

    add_heading(doc, "Profitability")
    add_table(doc,
        ["", "Year 1", "Year 2", "Year 3"],
        [
            ["Revenue", "TZS 30M", "TZS 180M", "TZS 720M"],
            ["Costs", "TZS 15M", "TZS 48M", "TZS 120M"],
            ["Net Profit", "TZS 15M", "TZS 132M", "TZS 600M"],
            ["Net Margin", "50%", "73%", "83%"],
            ["Break-even month", "Month 8", "-", "-"],
        ]
    )

    add_heading(doc, "Key Metrics")
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Customer Acquisition Cost (CAC)", "TZS 30,000 (~$12)"],
            ["Lifetime Value (LTV)", "TZS 1,200,000 (~$460)"],
            ["LTV:CAC Ratio", "40:1"],
            ["Payback Period", "< 1 month"],
            ["Gross Margin", "85%+"],
            ["Monthly churn rate", "5%"],
        ]
    )

    path = os.path.join(OUT_DIR, "Nyumba_Yangu_Financial_Projections.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# 4. MARKETING STRATEGY
# ============================================================
def create_marketing_strategy():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "MARKETING STRATEGY", "Go-To-Market Plan for Dar es Salaam")

    add_heading(doc, "1. Target Audience")
    add_heading(doc, "Primary: Landlords", level=2)
    doc.add_paragraph(
        "Who: Individuals owning 3-50 rental units in Dar es Salaam. "
        "Typically business owners, retirees, or employed professionals with rental income.\n\n"
        "Age: 35-60\n"
        "Tech level: Smartphone users (WhatsApp, M-Pesa) but not tech-savvy\n"
        "Pain: Tracking payments, managing tenants, filling vacancies\n"
        "Where to find them: At their properties, real estate events, WhatsApp groups, "
        "church and mosque communities"
    )

    add_heading(doc, "Secondary: Property Seekers", level=2)
    doc.add_paragraph(
        "Who: Young professionals and students looking for rentals in Dar es Salaam.\n\n"
        "Age: 20-40\n"
        "Tech level: Heavy smartphone users, active on Instagram/TikTok\n"
        "Pain: Broker fees, scam listings, no trusted platform\n"
        "Where to find them: Social media, universities, workplaces"
    )

    add_heading(doc, "2. Channels")

    add_heading(doc, "A. Direct Sales (Landlord Onboarding)", level=2)
    doc.add_paragraph(
        "This is our most important channel. Landlords need hands-on onboarding.\n\n"
        "How:\n"
        "- Field agents visit apartment buildings in target districts\n"
        "- Talk to caretakers/watchmen to identify the landlord\n"
        "- Demo the platform on a phone\n"
        "- Help them register their first property and units on the spot\n"
        "- Follow up within 48 hours\n\n"
        "Target districts (in order):\n"
        "1. Kinondoni - largest residential area\n"
        "2. Ilala - city center, mixed commercial/residential\n"
        "3. Ubungo - university area, high rental demand\n"
        "4. Temeke - growing residential area\n"
        "5. Kigamboni - new development area"
    )

    add_heading(doc, "B. WhatsApp Marketing", level=2)
    doc.add_paragraph(
        "WhatsApp is the primary communication tool in Tanzania.\n\n"
        "Strategy:\n"
        "- Join existing property/real estate WhatsApp groups\n"
        "- Share verified listings from our platform (with nyumbayangu.online links)\n"
        "- Create our own Nyumba Yangu community groups per district\n"
        "- Send broadcast messages to landlords about new features\n"
        "- Use WhatsApp Business for customer support"
    )

    add_heading(doc, "C. Social Media", level=2)
    doc.add_paragraph(
        "Instagram:\n"
        "- Property showcase posts (beautiful photos of listed properties)\n"
        "- Landlord success stories / testimonials\n"
        "- 'Did you know' posts about tenant rights and landlord best practices\n"
        "- Reels showing the platform in action\n\n"
        "TikTok:\n"
        "- Short videos: 'How to find a rental in DSM without a dalali'\n"
        "- Before/after: paper ledger vs. Nyumba Yangu dashboard\n"
        "- Apartment tour style content with listings\n\n"
        "Facebook:\n"
        "- Community group for landlords\n"
        "- Targeted ads to property owners in Dar es Salaam"
    )

    add_heading(doc, "D. Referral Program", level=2)
    doc.add_paragraph(
        "Landlord-to-landlord:\n"
        "- Existing landlord refers another landlord\n"
        "- Both get 1 month free subscription when the new landlord upgrades to paid\n"
        "- Simple referral code system\n\n"
        "Tenant-to-seeker:\n"
        "- Tenants share listings with friends looking for rentals\n"
        "- Organic growth through the marketplace"
    )

    add_heading(doc, "E. Partnerships", level=2)
    doc.add_paragraph(
        "Universities:\n"
        "- Partner with UDSM, Ardhi University, IFM for student housing listings\n"
        "- Posters on notice boards, mentions in student groups\n\n"
        "Real estate events:\n"
        "- Attend/sponsor local real estate exhibitions\n"
        "- Demo the platform to landlord associations\n\n"
        "Mobile money agents:\n"
        "- Partner with M-Pesa agents near residential areas\n"
        "- They recommend Nyumba Yangu to landlords who come for transactions"
    )

    add_heading(doc, "3. Content Strategy")
    doc.add_paragraph(
        "Weekly content calendar:\n\n"
        "Monday: New listing showcase (Instagram/WhatsApp)\n"
        "Tuesday: Landlord tip or feature highlight\n"
        "Wednesday: Market insight (rental prices by district)\n"
        "Thursday: Tenant tip (rights, payment best practices)\n"
        "Friday: Success story or testimonial\n"
        "Weekend: Apartment tour content (TikTok/Reels)"
    )

    add_heading(doc, "4. Budget (Monthly)")
    add_table(doc,
        ["Activity", "Monthly Budget (TZS)", "Notes"],
        [
            ["Field agent transport", "200,000", "Fuel/bus fare for property visits"],
            ["WhatsApp Business", "10,000", "Business account"],
            ["Instagram/Facebook ads", "150,000", "Targeted ads in DSM"],
            ["Printed materials (flyers)", "50,000", "For field visits"],
            ["Content creation", "100,000", "Photography, video editing"],
            ["Referral rewards", "100,000", "Free months for referrers"],
            ["TOTAL", "610,000", "~$235/month"],
        ]
    )

    add_heading(doc, "5. KPIs (Key Performance Indicators)")
    add_table(doc,
        ["Metric", "Monthly Target (Y1)"],
        [
            ["New landlord registrations", "15-20"],
            ["New paid subscriptions", "8-10"],
            ["New listings published", "30-50"],
            ["Website visits", "2,000+"],
            ["App downloads", "500+"],
            ["WhatsApp group members", "500+"],
            ["Instagram followers", "1,000+"],
            ["Landlord referrals", "5+"],
        ]
    )

    path = os.path.join(OUT_DIR, "Nyumba_Yangu_Marketing_Strategy.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# 5. COMPETITIVE ANALYSIS
# ============================================================
def create_competitive_analysis():
    doc = Document()
    set_style(doc)
    add_title_page(doc, "COMPETITIVE ANALYSIS", "Tanzania PropTech Landscape")

    add_heading(doc, "1. Market Overview")
    doc.add_paragraph(
        "Tanzania's property technology (PropTech) market is at a very early stage compared to "
        "neighboring Kenya (which has platforms like BuyRentKenya, Property24 Kenya) or Nigeria "
        "(PropertyPro, Estate Intel). Most property transactions in Tanzania still happen through "
        "informal channels: word-of-mouth, physical broker offices, and social media groups.\n\n"
        "This presents both a challenge (low digital adoption) and a massive opportunity "
        "(greenfield market with no dominant player)."
    )

    add_heading(doc, "2. Competitor Breakdown")

    add_heading(doc, "A. Informal Channels", level=2)
    add_heading(doc, "WhatsApp Groups & Social Media", level=3)
    add_table(doc,
        ["Aspect", "Detail"],
        [
            ["Description", "Property listings shared in WhatsApp groups, Instagram, Facebook"],
            ["Market share", "~60% of how properties are found"],
            ["Strengths", "Everyone already uses WhatsApp; zero learning curve; large reach"],
            ["Weaknesses", "No verification; scam listings common; listings disappear; no management tools; no payment tracking"],
            ["Threat to us", "Low - we complement rather than compete. Landlords can share our listings on WhatsApp."],
        ]
    )

    add_heading(doc, "Dalali (Brokers)", level=3)
    add_table(doc,
        ["Aspect", "Detail"],
        [
            ["Description", "Physical brokers who connect seekers with landlords for a fee"],
            ["Market share", "~30% of rental transactions"],
            ["Fee structure", "1-3 months' rent as commission"],
            ["Strengths", "Local knowledge; personal relationships; handle physical viewings"],
            ["Weaknesses", "Expensive; inconsistent service; no digital tools; some are unreliable"],
            ["Threat to us", "Medium - entrenched relationships, but young renters increasingly avoid them"],
        ]
    )

    add_heading(doc, "B. Digital Competitors", level=2)

    add_heading(doc, "Tanzania Property Centre", level=3)
    add_table(doc,
        ["Aspect", "Detail"],
        [
            ["Website", "tanzaniapropertycentre.com"],
            ["Type", "Listing portal"],
            ["Description", "Online property listing directory, mostly for sale properties"],
            ["Strengths", "Established brand; good SEO; large listing database"],
            ["Weaknesses", "No management tools; no tenant portal; no M-Pesa payment; no mobile app; mostly sale listings not rentals"],
            ["Differentiation", "We offer the full stack (marketplace + management + tenant portal). They are listings-only."],
        ]
    )

    add_heading(doc, "BeiYaUziki", level=3)
    add_table(doc,
        ["Aspect", "Detail"],
        [
            ["Type", "Listing site / classifieds"],
            ["Description", "General classifieds including some property listings"],
            ["Strengths", "Local brand; Swahili-first"],
            ["Weaknesses", "Not property-focused; basic listings only; no SaaS features"],
            ["Differentiation", "We are purpose-built for property management, not a classifieds site."],
        ]
    )

    add_heading(doc, "Zoom Tanzania", level=3)
    add_table(doc,
        ["Aspect", "Detail"],
        [
            ["Type", "Classifieds platform"],
            ["Description", "Online marketplace for various products including some property"],
            ["Strengths", "Broad user base"],
            ["Weaknesses", "Property is a small category; no specialization; no management tools"],
            ["Differentiation", "We specialize in property. Depth beats breadth."],
        ]
    )

    add_heading(doc, "3. Competitive Matrix")
    add_table(doc,
        ["Feature", "Nyumba Yangu", "TZ Property Centre", "Dalali", "WhatsApp"],
        [
            ["Property listings", "Yes", "Yes", "No", "Informal"],
            ["Verified listings", "Yes", "Partial", "No", "No"],
            ["Search & filter", "Yes", "Yes", "No", "No"],
            ["Landlord dashboard", "Yes", "No", "No", "No"],
            ["Tenant management", "Yes", "No", "No", "No"],
            ["Digital leases", "Yes", "No", "No", "No"],
            ["E-signature", "Yes", "No", "No", "No"],
            ["M-Pesa payments", "Yes", "No", "No", "No"],
            ["Payment tracking", "Yes", "No", "No", "No"],
            ["Maintenance requests", "Yes", "No", "No", "No"],
            ["Income reports", "Yes", "No", "No", "No"],
            ["Mobile app", "Yes", "No", "No", "Yes (chat)"],
            ["Amenity details", "Yes", "Basic", "Verbal", "No"],
            ["Free tier", "Yes", "Yes", "No", "Free"],
            ["Swahili support", "Yes (leases)", "Partial", "Yes", "Yes"],
            ["No broker fee", "Yes", "Yes", "No (1-3 months)", "Yes"],
        ]
    )

    add_heading(doc, "4. Our Competitive Advantages")
    doc.add_paragraph(
        "1. Full-stack platform: We are the only player in Tanzania combining a public marketplace "
        "with a landlord management suite and a tenant portal. Competitors offer one piece at most.\n\n"
        "2. Built for Tanzania: Swahili lease documents, M-Pesa integration, LUKU/DAWASA amenity "
        "tracking, Dar es Salaam district/ward structure. Not a generic template.\n\n"
        "3. Mobile-first: Our React Native app works on the phones Tanzanians actually use. "
        "Most competitors are desktop-only websites.\n\n"
        "4. Data moat: Once a landlord has tenants, leases, payment schedules, and history on our "
        "platform, switching to a competitor means losing all that data. This creates strong retention.\n\n"
        "5. Low cost structure: We built everything in-house with no expensive third-party dependencies. "
        "Our monthly server costs are under $20. This lets us offer a generous free tier while "
        "remaining profitable on paid tiers.\n\n"
        "6. First-mover in PropTech SaaS: While listing sites exist, no one in Tanzania offers "
        "property management as a service. We define the category."
    )

    add_heading(doc, "5. Risks from Competition")
    add_table(doc,
        ["Risk", "Likelihood", "Our Response"],
        [
            ["Kenya-based PropTech expands to Tanzania", "Medium", "We have local advantage: Swahili leases, DSM districts, M-Pesa TZ integration. Hard for outsiders to localize."],
            ["Tanzania Property Centre adds management tools", "Low", "They would need to rebuild from scratch. Our product is 2+ years ahead."],
            ["New local startup copies our model", "Medium", "First-mover advantage + data moat. Focus on execution speed and landlord relationships."],
            ["WhatsApp launches business tools for property", "Low", "WhatsApp is too generic for property management. We complement WhatsApp, not compete."],
        ]
    )

    add_heading(doc, "6. Conclusion")
    doc.add_paragraph(
        "The competitive landscape in Tanzania's PropTech market is remarkably empty. "
        "There is no established player offering what Nyumba Yangu offers. "
        "Our biggest competitor is the status quo: paper ledgers and WhatsApp. "
        "Our job is not to beat a competitor - it's to convince landlords that technology "
        "makes their life easier. Once they try it, the product sells itself."
    )

    path = os.path.join(OUT_DIR, "Nyumba_Yangu_Competitive_Analysis.docx")
    doc.save(path)
    print(f"Created: {path}")


# ============================================================
# RUN ALL
# ============================================================
if __name__ == "__main__":
    print("Generating Nyumba Yangu business documents...\n")
    create_business_plan()
    create_pitch_deck()
    create_financial_projections()
    create_marketing_strategy()
    create_competitive_analysis()
    print("\nAll documents generated successfully!")
